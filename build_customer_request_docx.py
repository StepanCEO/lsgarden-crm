from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, *, size=11, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, margin_twips="120"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "bottom", "start", "end"):
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), margin_twips)
        el.set(qn("w:type"), "dxa")


out = Path(r"C:\Users\User\Desktop\Projects\CRM\docs\customer_request_for_crm.docx")
out.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, side, Inches(1))
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)

for name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Запрос данных для запуска CRM")
set_run_font(r, size=24, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r = p.add_run("Ниже список данных и доступов, которые нужны, чтобы запустить CRM сразу полностью, без деления на этапы.")
set_run_font(r, size=11, color=RGBColor(70, 70, 70))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Если удобно, можно прислать всё одним сообщением или одним файлом по каждому блоку.")
set_run_font(r, size=10.5, italic=True, color=RGBColor(90, 90, 90))

rows = [
    ("VPS и запуск", "SSH-доступ к серверу, ОС, домен, доступ к DNS, информация по nginx, SSL и подтверждение, будем ли запускать через Docker."),
    ("S3", "Провайдер S3, bucket, access key / secret key, регион и что именно хранить: медиа, копии, бэкапы."),
    ("1C", "Способ интеграции: API, XML, CSV или другой вариант. Также нужна документация или тестовый доступ и список данных для обмена: товары, остатки, цены, заказы, клиенты."),
    ("Каналы связи", "VK, Avito, Telegram, почта и сайт: токены, доступы, SMTP-параметры, webhook или API, если они есть."),
    ("Пользователи и права", "Список сотрудников, роли, кто админ, и кто что должен видеть: клиентов, тикеты, задачи, аналитику, админку."),
    ("Бизнес-логика", "Статусы клиентов и тикетов, скрипты автоответов, правила уведомлений, кто получает задачи и как должны распределяться обращения."),
    ("Git и код", "Ссылка на репозиторий, доступы, кто принимает изменения и как удобнее вести работу: через один основной branch или через dev/main."),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = False
table.style = "Table Grid"

widths_dxa = [2520, 6840]
for idx, width in enumerate([Inches(1.75), Inches(4.75)]):
    table.columns[idx].width = width

header_cells = table.rows[0].cells
header_cells[0].text = "Блок"
header_cells[1].text = "Что нужно прислать"
for c in header_cells:
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = c._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    tc_pr.append(shd)
    set_cell_margins(c)
    for p in c.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size=10.5, bold=True)

for label, text in rows:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = text
    for idx, cell in enumerate(cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(widths_dxa[idx]))
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, size=10.5)

tbl = table._tbl
tbl_pr = tbl.tblPr
tbl_w = tbl_pr.first_child_found_in("w:tblW")
if tbl_w is None:
    tbl_w = OxmlElement("w:tblW")
    tbl_pr.append(tbl_w)
tbl_w.set(qn("w:type"), "dxa")
tbl_w.set(qn("w:w"), "9360")

tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
if tbl_ind is None:
    tbl_ind = OxmlElement("w:tblInd")
    tbl_pr.append(tbl_ind)
tbl_ind.set(qn("w:type"), "dxa")
tbl_ind.set(qn("w:w"), "120")

grid = tbl.tblGrid
for idx, width in enumerate(widths_dxa):
    grid.gridCol_lst[idx].w = width

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Если удобно, можно прислать всё одним сообщением.")
set_run_font(r, size=10.5, color=RGBColor(80, 80, 80))

for paragraph in doc.paragraphs:
    paragraph.paragraph_format.line_spacing = 1.15

doc.save(str(out))
print(out)
